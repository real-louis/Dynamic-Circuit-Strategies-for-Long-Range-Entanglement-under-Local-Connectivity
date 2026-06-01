#!/usr/bin/env python3
"""
Export an interactive-web friendly JSON from the English Word report.

Output is written to docs/content.json.
"""

from __future__ import annotations

import json
import re
from dataclasses import dataclass
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
REPORT = ROOT / "reports" / "Long_Range_Entanglement_Dynamic_Circuits_Report_EN.docx"
CFG = ROOT / "config" / "submission.json"
OUT = ROOT / "docs" / "content.json"


def _load_cfg() -> dict:
    if not CFG.is_file():
        return {}
    return json.loads(CFG.read_text(encoding="utf-8"))


def _clean(s: str) -> str:
    s = s.replace("\u00a0", " ").strip()
    s = re.sub(r"\s+", " ", s)
    return s


@dataclass
class Section:
    id: str
    title: str
    tagline: str | None
    paragraphs: list[str]


def _slugify(title: str) -> str:
    t = re.sub(r"[^a-zA-Z0-9]+", "-", title.strip()).strip("-").lower()
    return t or "section"


def extract_sections(doc: Document) -> list[Section]:
    sections: list[Section] = []
    cur: Section | None = None
    heading_re = re.compile(r"^\d+(\.\d+)*\.\s+.+")

    def flush() -> None:
        nonlocal cur
        if cur is None:
            return
        cur.paragraphs = [p for p in cur.paragraphs if p]
        if cur.paragraphs or cur.tagline:
            base = cur.id
            k = 2
            while any(s.id == cur.id for s in sections):
                cur.id = f"{base}-{k}"
                k += 1
            sections.append(cur)
        cur = None

    for p in doc.paragraphs:
        text = _clean(p.text)
        if not text:
            continue
        style = (p.style.name or "").lower() if p.style is not None else ""
        is_heading = "heading" in style or bool(heading_re.match(text))
        if is_heading:
            flush()
            sid = _slugify(text)
            cur = Section(id=sid, title=text, tagline=None, paragraphs=[])
            continue

        if cur is None:
            cur = Section(id="preface", title="Preface", tagline=None, paragraphs=[])

        if any(
            text.startswith(x)
            for x in (
                "PRX reproduce:",
                "PRX mechanism",
                "PRX qualitative analogy",
                "Ephys Challenge",
            )
        ):
            cur.tagline = text if cur.tagline is None else f"{cur.tagline} {text}"
        else:
            cur.paragraphs.append(text)

    flush()
    return sections


def extract_mapping_table(doc: Document) -> dict | None:
    for tbl in doc.tables:
        if not tbl.rows:
            continue
        headers = [_clean(c.text) for c in tbl.rows[0].cells]
        hset = {h.lower() for h in headers}
        if "paper location" in hset and "relation" in hset:
            rows = []
            for r in tbl.rows[1:]:
                rows.append([_clean(c.text) for c in r.cells])
            return {"headers": headers, "rows": rows}
    return None


def main() -> int:
    if not REPORT.is_file():
        raise FileNotFoundError(REPORT)

    cfg = _load_cfg()
    doc = Document(str(REPORT))

    gh = (cfg.get("github_repository_url") or "").rstrip("/")
    parts = gh.split("/")
    owner_repo = "/".join(parts[-2:]) if len(parts) >= 2 else ""
    raw_base = f"https://raw.githubusercontent.com/{owner_repo}/main/" if owner_repo else ""

    sections = extract_sections(doc)
    mapping = extract_mapping_table(doc)

    figures = [
        {"path": "figures/figure_chsh_comparison.png", "caption": "CHSH S value: noiseless baseline and noisy estimates"},
        {"path": "figures/crossover_simulation_with_ci.png", "caption": "Noisy crossover with 95% Wilson CI bands"},
        {"path": "figures/resources_cx_vs_L.png", "caption": "Ideal CNOT resource scaling vs L"},
        {"path": "figures/error_budget_bell_prep.png", "caption": "Paper error-budget bound (Bell preparation task)"},
        {"path": "figures/error_budget_lrcx.png", "caption": "Paper error-budget bound (LRCX task)"},
    ]

    out = {
        "title": "Dynamic Circuit Strategies for Long-Range Entanglement under Local Connectivity",
        "author": cfg.get("author_name_en", ""),
        "affiliation": cfg.get("affiliation_en", ""),
        "date": cfg.get("report_date", ""),
        "repository_commit": cfg.get("repository_commit", ""),
        "repo_raw_base": raw_base,
        "figures": figures,
        "tables": {"mapping": mapping},
        "sections": [
            {
                "id": s.id,
                "title": s.title,
                "tagline": s.tagline,
                "paragraphs": s.paragraphs[:18],
            }
            for s in sections
        ],
    }

    OUT.parent.mkdir(parents=True, exist_ok=True)
    OUT.write_text(json.dumps(out, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"Wrote {OUT}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())

