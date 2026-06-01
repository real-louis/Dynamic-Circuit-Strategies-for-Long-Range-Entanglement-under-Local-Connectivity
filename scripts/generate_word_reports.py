#!/usr/bin/env python3
"""Generate bilingual Word reports with embedded figures, CI tables, and real hyperlinks."""

from __future__ import annotations

import csv
import json
import math
import sys
from datetime import date
from pathlib import Path

import matplotlib.pyplot as plt
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

sys.path.insert(0, str(Path(__file__).resolve().parent))
from paths import ephys_root, repo_root
from report_figures import SHOTS, wilson_ci

ROOT = repo_root()
FIGURES = ROOT / "figures"
RESULTS = ROOT / "results"
REPORTS = ROOT / "reports"
CONFIG = ROOT / "config" / "submission.json"
EPHYS_SRC = ephys_root() / "src"


def load_config() -> dict:
    defaults = {
        "github_repository_url": "https://github.com/YOUR_USERNAME/long-range-entanglement-dynamic-circuits",
        "repository_commit": "",
        "author_name_en": "Jian-Yi He",
        "author_name_zh": "何建毅",
        "advisor_en": "Prof. Hsiu-Chuan Hsu",
        "advisor_zh": "許琇娟 教授",
        "affiliation_en": "Department of Applied Physics / Electronic Physics Program, National Chengchi University",
        "affiliation_zh": "國立政治大學 應用物理學系／電子物理學程",
        "contact_email": "louse2121@gmail.com",
        "report_date": date.today().strftime("%B %Y"),
    }
    if CONFIG.is_file():
        defaults.update(json.loads(CONFIG.read_text(encoding="utf-8")))
    return defaults


def load_crossover_rows() -> list[dict]:
    p = RESULTS / "crossover_and_resources.csv"
    if not p.is_file():
        raise FileNotFoundError(f"Missing {p}. Run: python3 scripts/build_graduate_package.py")
    with p.open(encoding="utf-8") as f:
        return list(csv.DictReader(f))


def load_formula_rows() -> list[dict]:
    with (RESULTS / "resource_formula_table.csv").open(encoding="utf-8") as f:
        return list(csv.DictReader(f))


def load_prx_cnot_rows() -> list[dict]:
    p = RESULTS / "prx_cnot_verification.csv"
    if not p.is_file():
        return []
    with p.open(encoding="utf-8") as f:
        return list(csv.DictReader(f))


def load_prx_ghz_rows() -> list[dict]:
    p = RESULTS / "prx_ghz_verification.csv"
    if not p.is_file():
        return []
    with p.open(encoding="utf-8") as f:
        return list(csv.DictReader(f))


def add_hyperlink(paragraph, text: str, url: str) -> None:
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run_el = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    r_pr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    run_el.append(r_pr)
    text_el = OxmlElement("w:t")
    text_el.text = text
    run_el.append(text_el)
    hyperlink.append(run_el)
    paragraph._p.append(hyperlink)


def _p(doc: Document, text: str, *, bold: bool = False, italic: bool = False, size: int = 11) -> None:
    r = doc.add_paragraph()
    run = r.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    r.paragraph_format.space_after = Pt(6)
    r.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def _h(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def _fig(doc: Document, path: Path, caption: str, width_in: float = 6.0) -> None:
    if not path.is_file():
        _p(doc, f"[Figure missing: {path.name}]", italic=True)
        return
    doc.add_picture(str(path), width=Inches(width_in))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    r.italic = True
    r.font.size = Pt(10)
    cap.paragraph_format.space_after = Pt(12)


def _code(doc: Document, title: str, code: str) -> None:
    h = doc.add_paragraph()
    ht = h.add_run(title)
    ht.bold = True
    ht.font.size = Pt(10)
    for line in code.rstrip("\n").split("\n"):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.22)
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(line if line.strip() else " ")
        run.font.name = "Courier New"
        run.font.size = Pt(8.5)
    doc.add_paragraph()


def _table(doc: Document, caption: str, headers: list[str], rows: list[list[str]]) -> None:
    cap = doc.add_paragraph()
    cr = cap.add_run(caption)
    cr.italic = True
    cr.font.size = Pt(10)
    cap.paragraph_format.space_after = Pt(6)
    if not rows:
        _p(doc, "[Table data not generated — run build_graduate_package.py]", italic=True)
        return
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, htxt in enumerate(headers):
        tbl.rows[0].cells[j].text = htxt
        for run in tbl.rows[0].cells[j].paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(9)
    for i, row in enumerate(rows, start=1):
        for j, cell in enumerate(row):
            tbl.rows[i].cells[j].text = cell
            for p in tbl.rows[i].cells[j].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
    doc.add_paragraph()


def _resource_row_at_L(form_rows: list[dict], L: int) -> list[list[str]]:
    row = next(r for r in form_rows if int(r["L"]) == L)
    out = [
        ["Static GHZ (single-rail)", row["static_cx"], row["static_depth"], "0"],
    ]
    if (row.get("swap_cx") or "").strip():
        out.append(["Entanglement swapping", row["swap_cx"], row["swap_depth"], str(max(0, L - 1))])
    out.append(["Two-rail zigzag", row["zigzag_cx"], row["zigzag_depth"], "0"])
    return out


def export_circuit_figures() -> dict[str, Path]:
    out: dict[str, Path] = {}
    REPORTS.mkdir(parents=True, exist_ok=True)
    if not EPHYS_SRC.is_dir():
        return out
    sys.path.insert(0, str(EPHYS_SRC))
    try:
        from challenge_common import generate_ladder_bell_circuit, generate_swapping_circuit

        specs = [
            ("static_L5", generate_ladder_bell_circuit(5), "Static GHZ chain (L=5)"),
            ("swapping_L3", generate_swapping_circuit(3, True), "Entanglement swapping (L=3)"),
        ]
        for key, qc, title in specs:
            path = REPORTS / f"figure_circuit_{key}.png"
            try:
                fig = qc.draw(output="mpl", fold=28, scale=0.65, plot_barriers=False)
                fig.suptitle(title, fontsize=11)
                fig.savefig(path, bbox_inches="tight", facecolor="white", dpi=175)
                plt.close(fig)
            except Exception:
                txt = str(qc.draw(output="text", fold=28))
                nlines = max(1, txt.count("\n") + 1)
                fig, ax = plt.subplots(figsize=(10.5, min(2.5 + 0.12 * nlines, 20)))
                ax.axis("off")
                ax.text(0.01, 0.98, txt, transform=ax.transAxes, fontsize=7, family="monospace", va="top")
                fig.suptitle(title, fontsize=11)
                fig.savefig(path, bbox_inches="tight", facecolor="white", dpi=160)
                plt.close(fig)
            out[key] = path
    except Exception as exc:
        print(f"Circuit figure export skipped: {exc}")
    return out


def build_document(
    lang: str,
    cfg: dict,
    crossover: list[dict],
    form_rows: list[dict],
    prx_cnot: list[dict],
    prx_ghz: list[dict],
    circuits: dict[str, Path],
) -> Document:
    en = lang == "en"
    doc = Document()
    doc.styles["Normal"].font.name = "Times New Roman"
    doc.styles["Normal"].font.size = Pt(11)

    gh = cfg["github_repository_url"]
    commit = (cfg.get("repository_commit") or "").strip()
    repo_name = gh.rstrip("/").split("/")[-1] or "repository"

    if en:
        title = "Dynamic Circuit Strategies for Long-Range Entanglement under Local Connectivity"
        subtitle = "From Ephys Challenge 2026 to PRX Quantum 5.030339"
        author, advisor, aff = cfg["author_name_en"], cfg["advisor_en"], cfg["affiliation_en"]
    else:
        title = "局域連接下長程糾纏的動態電路策略"
        subtitle = "從 Ephys Challenge 2026 到 PRX Quantum 5.030339"
        author, advisor, aff = cfg["author_name_zh"], cfg["advisor_zh"], cfg["affiliation_zh"]

    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = tp.add_run(title)
    tr.bold = True
    tr.font.size = Pt(16)
    sp = doc.add_paragraph()
    sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sp.add_run(subtitle)
    sr.font.size = Pt(12)
    sr.italic = True
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"{author}\n{advisor}\n{aff}\n{cfg['report_date']}").font.size = Pt(11)
    doc.add_page_break()

    # Abstract
    _h(doc, "Abstract" if en else "摘要", 1)
    if en:
        _p(
            doc,
            "We study long-range Bell-state preparation on a ladder topology under nearest-neighbor gates, "
            "motivated by Ephys Challenge 2026. Strategy A (static GHZ chain, 2L−1 CNOTs) is compared with "
            "Strategy B (entanglement swapping, L CNOTs, odd L) under one shared Qiskit Aer noise model "
            f"({SHOTS} shots). We report resource formulas, ideal CHSH verification (2√2), endpoint diagnostic "
            "P(00)+P(11) with 95% Wilson confidence intervals, and a noisy crossover curve. "
            "We independently reproduce Fig. 1a/2b circuits from PRX Quantum 5, 030339 (35 pytest checks). "
            "A paper error-budget analysis (Appendix B.2) is applied separately to Bell preparation and "
            "long-range CNOT tasks; we emphasize qualitative trend mapping—not calibrated equality with simulation.",
        )
        _p(doc, "Keywords: dynamic quantum circuits, long-range entanglement, entanglement swapping, NISQ", italic=True)
    else:
        _p(
            doc,
            "本報告研究梯子拓樸局域連接下的長程 Bell 態製備（Ephys Challenge 2026）。"
            f"比較策略甲（靜態 GHZ，2L−1 CNOT）與策略乙（糾纏交換，L CNOT），共用 Aer 雜訊模型（{SHOTS} shots）。"
            "含 CHSH 理想驗證、P(00)+P(11) 與 95% Wilson 信賴區間、PRX 5.030339 電路 reproduce（35 pytest）。"
            "error budget 分 Bell 製備與 LRCX 兩任務呈現，僅作定性對照，非與模擬定量等同。",
        )
        _p(doc, "關鍵詞：動態量子電路、長程糾纏、糾纏交換、NISQ", italic=True)

    if en:
        _h(doc, "Author contributions", 2)
        for item in [
            "Implemented and verified Ephys competition circuits with closed-form CX counts.",
            "Reproduced PRX Quantum 5.030339 LRCX/GHZ circuits with automated pytest equivalence.",
            "Unified qualitative crossover discussion via separate error-budget and simulation panels.",
        ]:
            doc.add_paragraph(item, style="List Bullet")

    # 1 Introduction
    _h(doc, "1. Introduction" if en else "1. 引言", 1)
    intro_en = (
        "Superconducting processors restrict two-qubit gates to nearest neighbors. Dynamic circuits use "
        "mid-circuit measurement and feed-forward for effectively constant-depth routing [Bäumer et al., "
        "PRX Quantum 5, 030339 (2024)]. We ask: under ladder topology, how do static versus measurement-based "
        "Bell preparation trade CX resources against noisy fidelity?"
    )
    intro_zh = (
        "超導處理器僅允許鄰接雙位元閘。動態電路以 measurement + feed-forward 達常數深度路由。"
        "問題：梯子拓樸下，靜態與測量輔助 Bell 製備如何權衡 CNOT 與含噪保真度？"
    )
    _p(doc, intro_en if en else intro_zh)

    # 2 Problem
    _h(doc, "2. Problem formulation" if en else "2. 問題形式化", 1)
    for b in (
        [
            "Single-rail abstraction: e0↔q0, e1↔qL, L segments, N=L+1 qubits.",
            "Two-rail zigzag uses 2(L+1) qubits on full ladder edges (comparison only).",
            "Target |Φ+⟩; metrics: CHSH, P(00)+P(11), CNOT count.",
        ]
        if en
        else [
            "單腳抽象：e0↔q0、e1↔qL，L 段，N=L+1。",
            "雙腳之字形用 2(L+1) 量子位元（對照策略）。",
            "目標 |Φ+⟩；指標：CHSH、P(00)+P(11)、CNOT。",
        ]
    ):
        doc.add_paragraph(b, style="List Bullet")

    # 3 Methods
    _h(doc, "3. Methods" if en else "3. 方法", 1)
    _h(doc, "3.1 Strategies A and B" if en else "3.1 策略甲、乙", 2)
    _p(
        doc,
        "Strategy A: H(e0) → forward CX → reverse CX (2L−1 CNOTs). "
        "Strategy B: adjacent |Φ+⟩ pairs, Bell measurement, X/Z feed-forward (L CNOTs, odd L). "
        "Noisy runs use classical post-decode when if_test is unstable in Aer."
        if en
        else "策略甲：H + 正/反向 CX（2L−1）。策略乙：|Φ+⟩ 對 + Bell 測量 + 前饋（L CNOT，L 奇）。含噪以古典解碼等價前饋。",
    )

    _h(doc, "3.2 Fair noise model" if en else "3.2 公平雜訊模型", 2)
    _table(
        doc,
        "Table 1. Shared Aer noise parameters." if en else "表 1. Aer 雜訊參數。",
        ["Parameter", "Value", "Note" if en else "說明"],
        [
            ["p(1q), p(2q)", "0.001, 0.01", "Depolarizing"],
            ["p_readout", "0.02", "Readout flip"],
            ["T1, T2", "60 μs, 80 μs", "Thermal relaxation"],
            ["Classical latency", "2500 ns", "Endpoint wait (swapping)"],
        ],
    )

    _h(doc, "3.3 PRX Quantum 5.030339 reproduction" if en else "3.3 PRX reproduce", 2)
    _code(doc, "Listing 1. PRX equivalence tests." if en else "程式 1. PRX 等價性測試。", "cd prx\npython3 -m pytest verify/test_equivalence.py -q  # 35 passed")

    # 4 Results
    _h(doc, "4. Results" if en else "4. 結果", 1)

    _h(doc, "4.1 Gate resources" if en else "4.1 資源", 2)
    _fig(doc, FIGURES / "resources_cx_vs_L.png", "Figure 1. Ideal CNOT count vs L." if en else "圖 1. 理想 CNOT vs L。")
    _table(
        doc,
        "Table 2. Resources at L=5 (from resource_formula_table.csv)." if en else "表 2. L=5 資源（CSV 自動讀取）。",
        ["Strategy", "CNOT", "Depth", "Mid-circuit meas."],
        _resource_row_at_L(form_rows, 5),
    )

    _h(doc, "4.2 CHSH verification" if en else "4.2 CHSH 驗證", 2)
    _fig(
        doc,
        FIGURES / "figure_chsh_comparison.png",
        "Figure 2. CHSH correlator: ideal (flat at 2√2) vs noisy static/zigzag." if en else "圖 2. CHSH：理想 vs 含噪。",
    )

    _h(doc, "4.3 Noisy Bell diagnostic" if en else "4.3 含噪 Bell 診斷", 2)
    end = [r for r in crossover if r.get("group") == "end_meas_A"]
    swap_map = {int(r["L"]): float(r["p_noisy_e0e1"]) for r in crossover if r.get("group") == "swap_decode_B"}
    sim_rows: list[list[str]] = []
    for strat, label in [("static_noisy", "Static"), ("zigzag_noisy", "Zigzag")]:
        for r in sorted([x for x in end if x["strategy"] == strat], key=lambda x: int(x["L"])):
            L = int(r["L"])
            p = float(r["p_noisy_e0e1"])
            lo, hi = wilson_ci(p, SHOTS)
            sim_rows.append([str(L), label, f"{p:.3f}", f"[{lo:.3f}, {hi:.3f}]"])
    for L, p in sorted(swap_map.items()):
        lo, hi = wilson_ci(p, SHOTS)
        sim_rows.append([str(L), "Swapping", f"{p:.3f}", f"[{lo:.3f}, {hi:.3f}]"])
    _table(
        doc,
        f"Table 3. P(00)+P(11) with 95% Wilson CI ({SHOTS} shots)." if en else f"表 3. P(00)+P(11) 與 95% CI（{SHOTS} shots）。",
        ["L", "Strategy", "Estimate", "95% CI"],
        sim_rows,
    )
    _fig(
        doc,
        FIGURES / "crossover_simulation_with_ci.png",
        "Figure 3. Noisy crossover with 95% CI (static, zigzag, swapping)." if en else "圖 3. 含噪 crossover（含 95% CI）。",
    )
    _p(
        doc,
        "Primary comparison: static vs swapping on single-rail endpoints. Zigzag (two-rail) uses more qubits/CNOT "
        "but respects full ladder edges; included as topology reference, not primary claim."
        if en
        else "主比較為單腳 static vs swapping；之字形為雙腳梯子對照，非主要結論。",
    )

    _h(doc, "4.4 PRX reproduction (quantitative)" if en else "4.4 PRX 定量驗證", 2)
    if prx_cnot:
        _table(
            doc,
            "Table 4. LRCX Bell fidelity (statevector, min over measurement branches)." if en else "表 4. LRCX Bell fidelity。",
            ["Ancilla n", "Dynamic min F", "Unitary SWAP F"],
            [[r["ancilla_n"], r["dynamic_bell_F"], r["unitary_swap_bell_F"]] for r in prx_cnot],
        )
    if prx_ghz:
        _table(
            doc,
            "Table 5. GHZ preparation fidelity." if en else "表 5. GHZ fidelity。",
            ["n", "Dynamic F", "Unitary F"],
            [[r["n_qubits"], r["ghz_dynamic_F"], r["ghz_unitary_F"]] for r in prx_ghz],
        )

    _h(doc, "4.5 Error-budget analysis (qualitative)" if en else "4.5 Error budget（定性）", 2)
    _p(
        doc,
        "Using Bäumer et al. Appendix B.2 parameters (μ≈3.65, λ_idle=0.03, λ_CNOT=0.02, λ_meas=0.03), "
        "F_proc ≥ exp(−λ_tot). These bounds use paper calibration—not our Aer model. Panels are split by task."
        if en
        else "依論文 Appendix B.2 參數計算 F_proc ≥ exp(−λ_tot)；與 Aer 參數不同，僅定性對照；分任務繪圖。",
    )
    _fig(
        doc,
        FIGURES / "error_budget_bell_prep.png",
        "Figure 4. Error budget: Bell-state preparation tasks." if en else "圖 4. Error budget：Bell 製備。",
    )
    _fig(
        doc,
        FIGURES / "error_budget_lrcx.png",
        "Figure 5. Error budget: long-range CNOT teleportation." if en else "圖 5. Error budget：LRCX。",
    )

    _h(doc, "4.6 Representative circuits" if en else "4.6 代表性電路", 2)
    if "static_L5" in circuits:
        _fig(doc, circuits["static_L5"], "Figure 6. Static chain (L=5)." if en else "圖 6. 靜態鏈 L=5。", 6.3)
    if "swapping_L3" in circuits:
        _fig(doc, circuits["swapping_L3"], "Figure 7. Swapping (L=3)." if en else "圖 7. 交換 L=3。", 6.3)

    _code(
        doc,
        "Listing 2. Reproduction." if en else "程式 2. 重現。",
        f"git clone {gh}\ncd {repo_name}\npip install -r requirements.txt\npython3 scripts/build_graduate_package.py",
    )

    # 5 Discussion
    _h(doc, "5. Discussion and limitations" if en else "5. 討論與限制", 1)
    for item in (
        [
            "Simulation uses Qiskit Aer; PRX Quantum reports 101-qubit IBM hardware.",
            "Error-budget curves and P(00)+P(11) use different noise calibrations—compare trends only.",
            "Bell-state preparation ≠ gate teleportation; tasks are related but not identical.",
            "Classical decode approximates feed-forward for fair noisy comparison.",
        ]
        if en
        else [
            "模擬為 Aer；論文為 IBM 實機。",
            "error budget 與 P(00)+P(11) 雜訊標定不同，僅比趨勢。",
            "Bell 製備 ≠ gate teleportation。",
            "含噪以古典解碼近似前饋。",
        ]
    ):
        doc.add_paragraph(item, style="List Number")

    # 6 Conclusion
    _h(doc, "6. Conclusion" if en else "6. 結論", 1)
    _p(
        doc,
        "We provide reproducible ladder-topology Bell preparation, PRX circuit reproduction with quantitative "
        "fidelities, and a carefully scoped crossover analysis separating simulation from paper error budgets."
        if en
        else "完成可重現 Bell 製備、PRX 定量 reproduce，以及分離模擬與 error budget 的 crossover 分析。",
    )

    if en:
        _h(doc, "Acknowledgments", 1)
        _p(doc, f"I thank {cfg['advisor_en']} for guidance on this project and the Ephys Challenge 2026 organizers.")

    # 7 Data availability
    _h(doc, "7. Data and code availability" if en else "7. 程式可用性", 1)
    _p(doc, "Source code and CSV artifacts:" if en else "原始碼與 CSV：")
    p = doc.add_paragraph()
    add_hyperlink(p, gh, gh)
    if commit and commit != "REPLACE_AFTER_GIT_PUSH":
        _p(doc, f"Reproduction reference commit: {commit}")
    else:
        _p(
            doc,
            "Set repository_commit in config/submission.json after git push, then regenerate this report."
            if en
            else "推送 GitHub 後更新 config/submission.json 的 repository_commit 並重產報告。",
            italic=True,
        )
    _p(doc, f"Contact: {cfg['contact_email']}", italic=True)

    _h(doc, "References" if en else "參考文獻", 1)
    for i, ref in enumerate(
        [
            "E. Bäumer et al., PRX Quantum 5, 030339 (2024). https://doi.org/10.1103/PRXQuantum.5.030339",
            "IBM Quantum, Long-range entanglement with dynamic circuits. https://docs.quantum.ibm.com/",
            "Ephys Challenge 2026, National Chengchi University.",
        ],
        1,
    ):
        doc.add_paragraph(f"[{i}] {ref}")

    doc.add_page_break()
    _h(doc, "Appendix A. Figure index" if en else "附錄 A. 圖表索引", 1)
    _table(
        doc,
        "Table A1." if en else "表 A1.",
        ["Fig.", "File", "Script"],
        [
            ["1", "resources_cx_vs_L.png", "ephys/competition_suite.py"],
            ["2", "figure_chsh_comparison.png", "scripts/report_figures.py"],
            ["3", "crossover_simulation_with_ci.png", "scripts/report_figures.py"],
            ["4–5", "error_budget_*.png", "scripts/error_budget_analysis.py"],
            ["6–7", "figure_circuit_*.png", "ephys/src/challenge_common.py"],
        ],
    )
    return doc


def main() -> None:
    REPORTS.mkdir(parents=True, exist_ok=True)
    cfg = load_config()
    crossover = load_crossover_rows()
    form_rows = load_formula_rows()
    prx_cnot = load_prx_cnot_rows()
    prx_ghz = load_prx_ghz_rows()
    circuits = export_circuit_figures()

    en_path = REPORTS / "Long_Range_Entanglement_Dynamic_Circuits_Report_EN.docx"
    build_document("en", cfg, crossover, form_rows, prx_cnot, prx_ghz, circuits).save(en_path)
    print(f"Wrote {en_path}")

    zh_path = REPORTS / "局域連接下長程糾纏動態電路策略_專題報告_ZH.docx"
    build_document("zh", cfg, crossover, form_rows, prx_cnot, prx_ghz, circuits).save(zh_path)
    print(f"Wrote {zh_path}")


if __name__ == "__main__":
    main()
